from __future__ import annotations

import os
import shutil
import sys
import tempfile
import threading
import traceback
import wave
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import numpy as np
try:
    import sounddevice as sd
except Exception:  # pragma: no cover
    sd = None

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document

try:
    from faster_whisper import WhisperModel
except Exception:  # pragma: no cover
    WhisperModel = None

APP_TITLE = "Audio to Text"
DEFAULT_MODEL_SIZE = os.getenv("WHISPER_MODEL_SIZE", "small.en")
DEFAULT_DEVICE = os.getenv("WHISPER_DEVICE", "cpu")
DEFAULT_COMPUTE_TYPE = os.getenv("WHISPER_COMPUTE_TYPE", "int8")
SAMPLE_RATE = 16000
CHANNELS = 1
SUPPORTED_FILE_TYPES = [
    ("Audio files", "*.wav *.mp3 *.m4a *.mp4 *.webm *.ogg *.flac *.aac"),
    ("WAV", "*.wav"),
    ("MP3", "*.mp3"),
    ("M4A", "*.m4a"),
    ("MP4", "*.mp4"),
    ("WEBM", "*.webm"),
    ("OGG", "*.ogg"),
    ("FLAC", "*.flac"),
    ("AAC", "*.aac"),
    ("All files", "*.*"),
]

BG_APP = "#F5F7FB"
BG_CARD = "#FFFFFF"
BG_NOTE = "#F8FAFD"
TEXT_MUTED = "#4B5563"
ACCENT = "#2563EB"
BORDER = "#D6DCE8"


def resource_path(relative_path: str) -> str:
    base_path = getattr(sys, "_MEIPASS", Path(__file__).resolve().parent)
    return str(Path(base_path) / relative_path)


@dataclass
class TranscriptResult:
    text: str
    language: str | None
    duration: float | None
    segments: list[dict[str, Any]]


class WhisperTranscriber:
    def __init__(self, model_size: str = DEFAULT_MODEL_SIZE) -> None:
        self.model_size = model_size
        self._model: Any | None = None
        self._lock = threading.Lock()

    def get_model(self) -> Any:
        if WhisperModel is None:
            raise RuntimeError("faster-whisper is not installed. Install requirements_desktop.txt before running the app.")
        with self._lock:
            if self._model is None:
                self._model = WhisperModel(
                    self.model_size,
                    device=DEFAULT_DEVICE,
                    compute_type=DEFAULT_COMPUTE_TYPE,
                )
            return self._model

    def transcribe(self, audio_path: str) -> TranscriptResult:
        model = self.get_model()
        segments, info = model.transcribe(
            audio_path,
            language="en",
            vad_filter=True,
            beam_size=5,
        )

        collected_segments: list[dict[str, Any]] = []
        full_text_parts: list[str] = []
        for seg in segments:
            seg_text = (seg.text or "").strip()
            collected_segments.append(
                {
                    "start": round(seg.start, 2),
                    "end": round(seg.end, 2),
                    "text": seg_text,
                }
            )
            if seg_text:
                full_text_parts.append(seg_text)

        transcript = " ".join(full_text_parts).strip()
        return TranscriptResult(
            text=transcript,
            language=getattr(info, "language", None),
            duration=getattr(info, "duration", None),
            segments=collected_segments,
        )


def save_pcm16_wav(output_path: str, audio: np.ndarray, sample_rate: int = SAMPLE_RATE) -> None:
    audio = np.asarray(audio, dtype=np.float32)
    if audio.ndim == 2 and audio.shape[1] == 1:
        audio = audio[:, 0]
    audio = np.clip(audio, -1.0, 1.0)
    pcm = (audio * 32767.0).astype(np.int16)
    with wave.open(output_path, "wb") as wav_file:
        wav_file.setnchannels(1)
        wav_file.setsampwidth(2)
        wav_file.setframerate(sample_rate)
        wav_file.writeframes(pcm.tobytes())


class AudioToTextApp:
    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self.root.title(APP_TITLE)
        self._fit_window()
        self.root.minsize(1080, 620)
        self.root.configure(bg=BG_APP)

        icon_path = Path(resource_path("app_icon.png"))
        if icon_path.exists():
            try:
                self.icon_image = tk.PhotoImage(file=str(icon_path))
                self.root.iconphoto(True, self.icon_image)
            except Exception:
                self.icon_image = None
        else:
            self.icon_image = None

        help_picture_path = Path(resource_path("Help_picture.png"))
        if help_picture_path.exists():
            try:
                self.help_image = tk.PhotoImage(file=str(help_picture_path))
            except Exception:
                self.help_image = None
        else:
            self.help_image = None

        self.transcriber = WhisperTranscriber(model_size=DEFAULT_MODEL_SIZE)
        self.current_audio_path: str | None = None
        self.current_audio_label: str = ""
        self.current_audio_origin: str = ""
        self.last_transcript: str = ""
        self.last_segments: list[dict[str, Any]] = []
        self.last_language: str | None = None
        self.last_duration: float | None = None
        self.last_timestamp: str = ""
        self.temp_recording_path: str | None = None

        self.recording_stream: Any | None = None
        self.recording_chunks: list[np.ndarray] = []
        self.is_recording = False
        self.is_playing = False

        self._build_ui()
        self._set_idle_status()
        self.root.bind("<Configure>", self._on_window_resize)
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    def _fit_window(self) -> None:
        self.root.update_idletasks()
        screen_w = self.root.winfo_screenwidth()
        screen_h = self.root.winfo_screenheight()
        width = min(1280, max(1120, screen_w - 80))
        height = min(690, max(620, screen_h - 120))
        pos_x = max((screen_w - width) // 2, 20)
        pos_y = max((screen_h - height) // 2 - 10, 20)
        self.root.geometry(f"{width}x{height}+{pos_x}+{pos_y}")

    def _build_ui(self) -> None:
        style = ttk.Style()
        try:
            style.theme_use("clam")
        except Exception:
            pass

        style.configure("App.TFrame", background=BG_APP)
        style.configure("Card.TFrame", background=BG_CARD)
        style.configure("Title.TLabel", background=BG_APP, font=("Segoe UI", 17, "bold"), foreground="#111827")
        style.configure("Sub.TLabel", background=BG_APP, font=("Segoe UI", 10), foreground=TEXT_MUTED)
        style.configure("Section.TLabel", background=BG_CARD, font=("Segoe UI", 10, "bold"), foreground="#111827")
        style.configure("Action.TButton", font=("Segoe UI", 10), padding=(10, 7))
        style.configure("TLabelframe", background=BG_CARD, bordercolor=BORDER, relief="solid")
        style.configure("TLabelframe.Label", background=BG_CARD, font=("Segoe UI", 10, "bold"), foreground="#111827")
        style.configure("Treeview", rowheight=24, font=("Segoe UI", 9))
        style.configure("Treeview.Heading", font=("Segoe UI", 9, "bold"))

        outer = ttk.Frame(self.root, padding=10, style="App.TFrame")
        outer.pack(fill="both", expand=True)
        outer.columnconfigure(0, weight=1)
        outer.rowconfigure(1, weight=1)

        header = ttk.Frame(outer, style="App.TFrame")
        header.grid(row=0, column=0, sticky="ew", pady=(0, 8))
        header.columnconfigure(0, weight=1)
        header.columnconfigure(1, weight=0)

        title_area = ttk.Frame(header, style="App.TFrame")
        title_area.grid(row=0, column=0, sticky="nw")
        if self.icon_image is not None:
            try:
                self.header_icon = self.icon_image.subsample(6, 6)
            except Exception:
                self.header_icon = self.icon_image
            ttk.Label(title_area, image=self.header_icon, style="Sub.TLabel").grid(row=0, column=0, rowspan=2, sticky="nw", padx=(0, 10))
        ttk.Label(title_area, text="Audio to Text", style="Title.TLabel").grid(row=0, column=1, sticky="w", pady=(2, 0))
        ttk.Label(
            title_area,
            text="Desktop audio transcription using microphone recording or audio-file upload.",
            style="Sub.TLabel",
            wraplength=560,
        ).grid(row=1, column=1, sticky="w", pady=(4, 0))

        right_header = ttk.Frame(header, style="App.TFrame")
        right_header.grid(row=0, column=1, sticky="ne", padx=(12, 0))
        right_header.columnconfigure(0, weight=1)
        right_header.columnconfigure(1, weight=0)

        disclaimer_frame = tk.Frame(
            right_header,
            bg=BG_NOTE,
            highlightbackground=BORDER,
            highlightthickness=1,
            padx=10,
            pady=7,
        )
        disclaimer_frame.grid(row=0, column=0, sticky="e")
        disclaimer_text = "Note: This code is created with the help of OpenAI\ncontact: arkabhowmik@yahoo.co.uk"
        disclaimer_label = tk.Label(
            disclaimer_frame,
            text=disclaimer_text,
            font=("Segoe UI", 9, "italic"),
            justify="center",
            bg=BG_NOTE,
            fg="#1F2937",
            anchor="center",
        )
        disclaimer_label.pack()

        help_holder = ttk.Frame(right_header, style="App.TFrame")
        help_holder.grid(row=0, column=1, sticky="ne", padx=(8, 0))
        if self.help_image is not None:
            self.help_widget = tk.Label(
                help_holder,
                image=self._make_help_image(),
                cursor="hand2",
                bg=BG_APP,
                bd=0,
                highlightthickness=0,
            )
            self.help_widget.pack(anchor="ne")
            self.help_widget.bind("<Button-1>", lambda _event: self.download_help_pdf())
        else:
            ttk.Button(help_holder, text="Help", command=self.download_help_pdf, style="Action.TButton").pack(anchor="ne")

        content = ttk.Frame(outer, style="App.TFrame")
        content.grid(row=1, column=0, sticky="nsew")
        content.columnconfigure(0, weight=0)
        content.columnconfigure(1, weight=1)
        content.rowconfigure(0, weight=1)

        left = ttk.Frame(content, width=300, style="App.TFrame")
        left.grid(row=0, column=0, sticky="nsw", padx=(0, 10))
        left.grid_propagate(False)

        right = ttk.Frame(content, style="App.TFrame")
        right.grid(row=0, column=1, sticky="nsew")
        right.columnconfigure(0, weight=1)
        right.rowconfigure(1, weight=3)
        right.rowconfigure(4, weight=2)

        input_card = ttk.LabelFrame(left, text="Input", padding=10)
        input_card.pack(fill="x", pady=(0, 8))

        self.file_label_var = tk.StringVar(value="No audio selected")
        ttk.Button(input_card, text="Open audio file", command=self.select_audio_file, style="Action.TButton").pack(fill="x")
        ttk.Label(input_card, textvariable=self.file_label_var, wraplength=250, background=BG_CARD, foreground=TEXT_MUTED).pack(anchor="w", pady=(8, 8))

        ttk.Separator(input_card, orient="horizontal").pack(fill="x", pady=4)
        ttk.Label(input_card, text="Microphone", style="Section.TLabel").pack(anchor="w", pady=(4, 8))
        self.record_btn = ttk.Button(input_card, text="Start recording", command=self.start_recording, style="Action.TButton")
        self.record_btn.pack(fill="x")
        self.stop_btn = ttk.Button(input_card, text="Stop recording", command=self.stop_recording, state="disabled", style="Action.TButton")
        self.stop_btn.pack(fill="x", pady=(6, 0))
        self.play_btn = ttk.Button(input_card, text="Play recording", command=self.play_recording, state="disabled", style="Action.TButton")
        self.play_btn.pack(fill="x", pady=(6, 0))

        action_card = ttk.LabelFrame(left, text="Actions", padding=10)
        action_card.pack(fill="x", pady=(0, 8))
        self.transcribe_btn = ttk.Button(action_card, text="Transcribe", command=self.transcribe_current_audio, style="Action.TButton")
        self.transcribe_btn.pack(fill="x")
        ttk.Button(action_card, text="Clear output", command=self.clear_output, style="Action.TButton").pack(fill="x", pady=(6, 0))

        info_card = ttk.LabelFrame(left, text="Status", padding=10)
        info_card.pack(fill="x")
        self.status_var = tk.StringVar(value="Ready")
        self.meta_var = tk.StringVar(value=f"Model: {DEFAULT_MODEL_SIZE} | Device: {DEFAULT_DEVICE} | Compute: {DEFAULT_COMPUTE_TYPE}")
        ttk.Label(info_card, textvariable=self.status_var, wraplength=250, background=BG_CARD).pack(anchor="w")
        ttk.Label(info_card, textvariable=self.meta_var, wraplength=250, background=BG_CARD, foreground=TEXT_MUTED).pack(anchor="w", pady=(8, 0))

        export_bar = ttk.Frame(right, style="App.TFrame")
        export_bar.grid(row=0, column=0, sticky="ew", pady=(0, 6))
        ttk.Button(export_bar, text="Download TXT", command=self.save_txt, style="Action.TButton").pack(side="left")
        ttk.Button(export_bar, text="Download DOCX", command=self.save_docx, style="Action.TButton").pack(side="left", padx=(8, 0))

        transcript_card = ttk.LabelFrame(right, text="Transcript", padding=8)
        transcript_card.grid(row=1, column=0, sticky="nsew")
        transcript_card.rowconfigure(0, weight=1)
        transcript_card.columnconfigure(0, weight=1)

        self.transcript_text = tk.Text(
            transcript_card,
            wrap="word",
            font=("Segoe UI", 10),
            height=10,
            bd=0,
            relief="flat",
            padx=8,
            pady=8,
            bg="#FFFFFF",
            fg="#111827",
            insertbackground="#111827",
        )
        self.transcript_text.grid(row=0, column=0, sticky="nsew")
        transcript_scroll = ttk.Scrollbar(transcript_card, orient="vertical", command=self.transcript_text.yview)
        transcript_scroll.grid(row=0, column=1, sticky="ns")
        self.transcript_text.configure(yscrollcommand=transcript_scroll.set)

        stats_frame = ttk.Frame(right, style="App.TFrame")
        stats_frame.grid(row=2, column=0, sticky="ew", pady=6)
        self.summary_var = tk.StringVar(value="Words: 0 | Segments: 0 | Source: N/A | Last run: N/A")
        ttk.Label(stats_frame, textvariable=self.summary_var, style="Sub.TLabel").pack(anchor="w")

        ttk.Label(right, text="Time-stamped Segments", style="Sub.TLabel").grid(row=3, column=0, sticky="w", pady=(0, 4))

        segments_card = ttk.LabelFrame(right, text="", padding=4)
        segments_card.grid(row=4, column=0, sticky="nsew")
        segments_card.rowconfigure(0, weight=1)
        segments_card.columnconfigure(0, weight=1)

        self.segment_tree = ttk.Treeview(segments_card, columns=("start", "end", "text"), show="headings", height=7)
        self.segment_tree.heading("start", text="Start (s)")
        self.segment_tree.heading("end", text="End (s)")
        self.segment_tree.heading("text", text="Text")
        self.segment_tree.column("start", width=110, anchor="center", stretch=False)
        self.segment_tree.column("end", width=110, anchor="center", stretch=False)
        self.segment_tree.column("text", width=650, anchor="w", stretch=True)
        self.segment_tree.grid(row=0, column=0, sticky="nsew")
        segment_scroll = ttk.Scrollbar(segments_card, orient="vertical", command=self.segment_tree.yview)
        segment_scroll.grid(row=0, column=1, sticky="ns")
        self.segment_tree.configure(yscrollcommand=segment_scroll.set)

        self._adjust_columns()

    def _make_help_image(self) -> tk.PhotoImage:
        if self.help_image is None:
            raise RuntimeError("Help image is not available.")
        try:
            width = self.help_image.width()
            if width >= 120:
                self.help_display_image = self.help_image.subsample(2, 2)
            else:
                self.help_display_image = self.help_image
        except Exception:
            self.help_display_image = self.help_image
        return self.help_display_image

    def _set_idle_status(self) -> None:
        self.status_var.set("Ready")

    def _on_window_resize(self, _event: Any = None) -> None:
        self._adjust_columns()

    def _adjust_columns(self) -> None:
        total_width = max(self.root.winfo_width(), 1080)
        text_width = max(420, total_width - 620)
        if hasattr(self, "segment_tree"):
            self.segment_tree.column("text", width=text_width)

    def set_status(self, text: str) -> None:
        self.status_var.set(text)
        self.root.update_idletasks()

    def set_current_audio(self, path: str, origin: str) -> None:
        self.current_audio_path = path
        self.current_audio_origin = origin
        self.current_audio_label = Path(path).name
        self.file_label_var.set(f"Selected: {self.current_audio_label}")
        self.summary_var.set(
            f"Words: {len(self.last_transcript.split())} | Segments: {len(self.last_segments)} | Source: {origin} | Last run: {self.last_timestamp or 'N/A'}"
        )
        self._update_play_button_state()

    def _update_play_button_state(self) -> None:
        has_recording = bool(self.temp_recording_path and Path(self.temp_recording_path).exists())
        if self.is_recording:
            self.play_btn.configure(state="disabled")
        elif has_recording:
            self.play_btn.configure(state="normal")
        else:
            self.play_btn.configure(state="disabled")

    def _cleanup_temp_recording_file(self) -> None:
        if self.temp_recording_path and Path(self.temp_recording_path).exists():
            try:
                Path(self.temp_recording_path).unlink()
            except Exception:
                pass
        self.temp_recording_path = None

    def _clear_selected_audio(self, delete_temp_recording: bool = True) -> None:
        if delete_temp_recording:
            self._cleanup_temp_recording_file()
        self.current_audio_path = None
        self.current_audio_origin = ""
        self.current_audio_label = ""
        self.file_label_var.set("No audio selected")
        self.recording_chunks = []
        self._update_play_button_state()

    def _reset_recording_state(self, discard_audio: bool = False) -> None:
        self._stop_playback()
        try:
            if self.recording_stream is not None:
                self.recording_stream.stop()
                self.recording_stream.close()
        except Exception:
            pass
        finally:
            self.recording_stream = None
            self.is_recording = False

        if discard_audio:
            self.recording_chunks = []

        if hasattr(self, "record_btn"):
            self.record_btn.configure(state="normal")
        if hasattr(self, "stop_btn"):
            self.stop_btn.configure(state="disabled")
        self._update_play_button_state()

    def select_audio_file(self) -> None:
        file_path = filedialog.askopenfilename(title="Select audio file", filetypes=SUPPORTED_FILE_TYPES)
        if not file_path:
            return
        self._cleanup_temp_recording_file()
        self.recording_chunks = []
        self.set_current_audio(file_path, "uploaded audio")
        self.set_status("Audio file selected. Click Transcribe to continue.")

    def start_recording(self) -> None:
        if sd is None:
            messagebox.showerror("Missing dependency", "sounddevice is not installed. Install requirements_desktop.txt before using microphone recording.")
            return
        if self.is_recording:
            return

        self._reset_recording_state(discard_audio=True)
        self.recording_chunks = []
        try:
            self.recording_stream = sd.InputStream(
                samplerate=SAMPLE_RATE,
                channels=CHANNELS,
                dtype="float32",
                callback=self._recording_callback,
            )
            self.recording_stream.start()
        except Exception as exc:
            messagebox.showerror("Microphone error", f"Could not start recording.\n\n{exc}")
            self.recording_stream = None
            self.recording_chunks = []
            return

        self.is_recording = True
        self.record_btn.configure(state="disabled")
        self.stop_btn.configure(state="normal")
        self.play_btn.configure(state="disabled")
        self.set_status("Recording from microphone...")

    def _recording_callback(self, indata: np.ndarray, frames: int, time_info: Any, status: Any) -> None:
        if status:
            pass
        self.recording_chunks.append(indata.copy())

    def stop_recording(self) -> None:
        if not self.is_recording:
            return

        self._reset_recording_state(discard_audio=False)

        if not self.recording_chunks:
            self._update_play_button_state()
            self.set_status("Recording stopped, but no audio was captured.")
            messagebox.showwarning("Recording", "No audio was captured from the microphone.")
            return

        audio = np.concatenate(self.recording_chunks, axis=0)
        temp_fd, temp_name = tempfile.mkstemp(prefix="audio_to_text_recording_", suffix=".wav")
        os.close(temp_fd)
        temp_path = Path(temp_name)
        save_pcm16_wav(str(temp_path), audio, SAMPLE_RATE)

        self._cleanup_temp_recording_file()
        self.temp_recording_path = str(temp_path)
        self.set_current_audio(str(temp_path), "recorded audio")
        self.set_status("Recording captured. Click Transcribe to continue.")

    def play_recording(self) -> None:
        if sd is None:
            messagebox.showerror("Missing dependency", "sounddevice is not installed. Install requirements_desktop.txt before playback.")
            return
        if not self.temp_recording_path or not Path(self.temp_recording_path).exists():
            messagebox.showwarning("No recording", "No recorded microphone audio is available to play.")
            return
        try:
            with wave.open(self.temp_recording_path, "rb") as wav_file:
                frames = wav_file.readframes(wav_file.getnframes())
                sample_rate = wav_file.getframerate()
                channels = wav_file.getnchannels()
            audio = np.frombuffer(frames, dtype=np.int16).astype(np.float32) / 32767.0
            if channels > 1:
                audio = audio.reshape(-1, channels)
            self._stop_playback()
            sd.play(audio, samplerate=sample_rate)
            self.is_playing = True
            self.set_status("Playing recorded audio...")
            threading.Thread(target=self._watch_playback, daemon=True).start()
        except Exception as exc:
            messagebox.showerror("Playback error", f"Could not play recorded audio.\n\n{exc}")

    def _watch_playback(self) -> None:
        try:
            if sd is not None:
                sd.wait()
        except Exception:
            pass
        self.root.after(0, self._finish_playback)

    def _finish_playback(self) -> None:
        if self.is_playing:
            self.is_playing = False
            self.set_status("Playback finished.")
            self._update_play_button_state()

    def _stop_playback(self) -> None:
        if sd is not None:
            try:
                sd.stop()
            except Exception:
                pass
        self.is_playing = False

    def download_help_pdf(self) -> None:
        help_pdf_path = Path(resource_path("help_audio_to_text.pdf"))
        if not help_pdf_path.exists():
            messagebox.showerror("Help file missing", "The bundled help PDF was not found.")
            return
        output_path = filedialog.asksaveasfilename(
            title="Download help PDF",
            defaultextension=".pdf",
            filetypes=[("PDF file", "*.pdf")],
            initialfile="help_audio_to_text.pdf",
        )
        if not output_path:
            return
        shutil.copyfile(help_pdf_path, output_path)
        self.set_status(f"Help PDF downloaded: {Path(output_path).name}")

    def transcribe_current_audio(self) -> None:
        if not self.current_audio_path or not Path(self.current_audio_path).exists():
            messagebox.showwarning("No audio", "Please open an audio file or record audio first.")
            return

        self._stop_playback()
        self.transcribe_btn.configure(state="disabled")
        self.record_btn.configure(state="disabled")
        self.stop_btn.configure(state="disabled")
        self.play_btn.configure(state="disabled")
        self.set_status("Loading model and transcribing audio...")

        worker = threading.Thread(target=self._transcribe_worker, daemon=True)
        worker.start()

    def _transcribe_worker(self) -> None:
        try:
            result = self.transcriber.transcribe(self.current_audio_path)
            self.root.after(0, lambda: self._apply_transcript_result(result))
        except Exception as exc:
            error_text = f"Transcription failed: {exc}"
            debug_trace = traceback.format_exc(limit=3)
            self.root.after(0, lambda: self._handle_transcription_error(error_text, debug_trace))

    def _apply_transcript_result(self, result: TranscriptResult) -> None:
        self.last_transcript = result.text
        self.last_segments = result.segments
        self.last_language = result.language
        self.last_duration = result.duration
        self.last_timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        self.transcript_text.delete("1.0", tk.END)
        self.transcript_text.insert("1.0", self.last_transcript)

        for item_id in self.segment_tree.get_children():
            self.segment_tree.delete(item_id)
        for seg in self.last_segments:
            self.segment_tree.insert(
                "",
                tk.END,
                values=(f"{seg.get('start', 0):.2f}", f"{seg.get('end', 0):.2f}", seg.get("text", "")),
            )

        word_count = len(self.last_transcript.split()) if self.last_transcript else 0
        source = self.current_audio_origin or "audio"
        self.summary_var.set(
            f"Words: {word_count} | Segments: {len(self.last_segments)} | Source: {source} | Last run: {self.last_timestamp}"
        )
        self.set_status("Transcription complete.")
        self._restore_actions_after_job()

    def _handle_transcription_error(self, error_text: str, debug_trace: str) -> None:
        self.set_status(error_text)
        self._restore_actions_after_job()
        messagebox.showerror("Transcription failed", f"{error_text}\n\n{debug_trace}")

    def _restore_actions_after_job(self) -> None:
        self.transcribe_btn.configure(state="normal")
        if self.is_recording:
            self.record_btn.configure(state="disabled")
            self.stop_btn.configure(state="normal")
            self.play_btn.configure(state="disabled")
        else:
            self.record_btn.configure(state="normal")
            self.stop_btn.configure(state="disabled")
            self._update_play_button_state()

    def clear_output(self) -> None:
        self._reset_recording_state(discard_audio=True)
        self._clear_selected_audio(delete_temp_recording=True)
        self.last_transcript = ""
        self.last_segments = []
        self.last_language = None
        self.last_duration = None
        self.last_timestamp = ""
        self.transcript_text.delete("1.0", tk.END)
        for item_id in self.segment_tree.get_children():
            self.segment_tree.delete(item_id)
        self.summary_var.set("Words: 0 | Segments: 0 | Source: N/A | Last run: N/A")
        self.set_status("Input and output cleared.")

    def save_txt(self) -> None:
        if not self.last_transcript:
            messagebox.showwarning("Nothing to save", "No transcript is available to save.")
            return
        initial_name = f"{self._export_stem()}_transcript.txt"
        output_path = filedialog.asksaveasfilename(
            title="Save TXT transcript",
            defaultextension=".txt",
            filetypes=[("Text file", "*.txt")],
            initialfile=initial_name,
        )
        if not output_path:
            return
        Path(output_path).write_text(self.last_transcript, encoding="utf-8")
        self.set_status(f"Saved TXT: {Path(output_path).name}")

    def save_docx(self) -> None:
        if not self.last_transcript:
            messagebox.showwarning("Nothing to save", "No transcript is available to save.")
            return
        initial_name = f"{self._export_stem()}_transcript.docx"
        output_path = filedialog.asksaveasfilename(
            title="Save DOCX transcript",
            defaultextension=".docx",
            filetypes=[("Word document", "*.docx")],
            initialfile=initial_name,
        )
        if not output_path:
            return

        document = Document()
        document.add_heading("Audio Transcript", level=1)
        meta = document.add_paragraph()
        meta.add_run("Source: ").bold = True
        meta.add_run(self.current_audio_origin or "audio")
        meta.add_run("\nFilename: ").bold = True
        meta.add_run(Path(self.current_audio_path).name if self.current_audio_path else "N/A")
        meta.add_run("\nGenerated: ").bold = True
        meta.add_run(self.last_timestamp or "N/A")
        meta.add_run("\nLanguage: ").bold = True
        meta.add_run(self.last_language or "N/A")
        meta.add_run("\nDuration (s): ").bold = True
        meta.add_run(str(self.last_duration) if self.last_duration is not None else "N/A")

        document.add_heading("Transcript", level=2)
        document.add_paragraph(self.last_transcript)

        if self.last_segments:
            document.add_heading("Time-stamped Segments", level=2)
            for seg in self.last_segments:
                document.add_paragraph(
                    f"[{seg.get('start', 0):.2f}s → {seg.get('end', 0):.2f}s] {seg.get('text', '')}"
                )

        document.save(output_path)
        self.set_status(f"Saved DOCX: {Path(output_path).name}")

    def _export_stem(self) -> str:
        if not self.current_audio_path:
            return "audio"
        stem = Path(self.current_audio_path).stem.strip()
        return stem or "audio"

    def on_close(self) -> None:
        self._reset_recording_state(discard_audio=True)
        self._cleanup_temp_recording_file()
        self.root.destroy()


def run_self_test() -> int:
    """Light sanity check that avoids GUI launch and model download."""
    tmp_dir = Path(tempfile.gettempdir())
    wav_path = tmp_dir / "audio_to_text_self_test.wav"
    txt_path = tmp_dir / "audio_to_text_self_test.txt"
    docx_path = tmp_dir / "audio_to_text_self_test.docx"
    help_copy_path = tmp_dir / "audio_to_text_help_copy.pdf"

    tone_seconds = 1.0
    timeline = np.linspace(0, tone_seconds, int(SAMPLE_RATE * tone_seconds), endpoint=False)
    tone = 0.1 * np.sin(2 * np.pi * 440 * timeline).astype(np.float32)
    save_pcm16_wav(str(wav_path), tone, SAMPLE_RATE)

    if not wav_path.exists() or wav_path.stat().st_size == 0:
        raise RuntimeError("WAV sanity check failed: file was not created.")

    txt_path.write_text("sanity check", encoding="utf-8")
    if txt_path.read_text(encoding="utf-8") != "sanity check":
        raise RuntimeError("TXT sanity check failed.")

    document = Document()
    document.add_heading("Audio Transcript", level=1)
    document.add_paragraph("sanity check")
    document.save(docx_path)
    if not docx_path.exists() or docx_path.stat().st_size == 0:
        raise RuntimeError("DOCX sanity check failed.")

    help_pdf_path = Path(resource_path("help_audio_to_text.pdf"))
    if not help_pdf_path.exists() or help_pdf_path.stat().st_size == 0:
        raise RuntimeError("Help PDF sanity check failed: bundled help file is missing.")
    shutil.copyfile(help_pdf_path, help_copy_path)
    if not help_copy_path.exists() or help_copy_path.stat().st_size == 0:
        raise RuntimeError("Help PDF copy sanity check failed.")

    print("Self-test passed.")
    return 0


def main() -> int:
    if "--self-test" in sys.argv:
        return run_self_test()

    root = tk.Tk()
    app = AudioToTextApp(root)
    root.mainloop()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
