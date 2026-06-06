from __future__ import annotations

import base64
import difflib
import importlib.util
import html
import io
import os
import tempfile
import threading
import time
import wave
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    import numpy as np
except Exception:  # pragma: no cover
    np = None

import pandas as pd
import streamlit as st
from docx import Document

try:
    import soundcard as sc
except Exception:  # pragma: no cover
    sc = None

try:
    from scipy import signal
except Exception:  # pragma: no cover
    signal = None

try:
    from faster_whisper import WhisperModel
except Exception:  # pragma: no cover
    WhisperModel = None

try:
    from streamlit_mic_recorder import mic_recorder
except Exception:  # pragma: no cover
    mic_recorder = None

def env_bool(name: str, default: bool) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


def env_int(name: str, default: int) -> int:
    try:
        return int(os.getenv(name, str(default)))
    except ValueError:
        return default


def env_float(name: str, default: float) -> float:
    try:
        return float(os.getenv(name, str(default)))
    except ValueError:
        return default


def module_available(name: str) -> bool:
    try:
        return importlib.util.find_spec(name) is not None
    except Exception:
        return False



# -----------------------------
# App configuration from environment variables
# Change these values through environment variables instead of editing code.
# -----------------------------
APP_TITLE = "Audio to Text"
DEFAULT_MODEL_SIZE = os.getenv("WHISPER_MODEL", os.getenv("WHISPER_MODEL_SIZE", "small.en"))
DEFAULT_DEVICE = os.getenv("WHISPER_DEVICE", "cpu")
DEFAULT_COMPUTE_TYPE = os.getenv("WHISPER_COMPUTE_TYPE", "int8")
WHISPER_LANGUAGE = os.getenv("WHISPER_LANGUAGE", "en") or None
WHISPER_BEAM_SIZE = env_int("WHISPER_BEAM_SIZE", 5)
WHISPER_BEST_OF = env_int("WHISPER_BEST_OF", 5)
WHISPER_VAD_FILTER = env_bool("WHISPER_VAD_FILTER", True)
WHISPER_NO_SPEECH_THRESHOLD = env_float("WHISPER_NO_SPEECH_THRESHOLD", 0.25)
WHISPER_LOG_PROB_THRESHOLD = env_float("WHISPER_LOG_PROB_THRESHOLD", -1.2)
WHISPER_CONDITION_ON_PREVIOUS_TEXT = env_bool("WHISPER_CONDITION_ON_PREVIOUS_TEXT", False)
WHISPER_TEMPERATURE = env_float("WHISPER_TEMPERATURE", 0.0)
TRANSCRIPTION_BACKEND = os.getenv("TRANSCRIPTION_BACKEND", "faster_whisper").strip().lower()
ASR_QUALITY_PRESET = os.getenv("ASR_QUALITY_PRESET", "balanced").strip().lower()
TRANSCRIBE_SOURCES_SEPARATELY = env_bool("TRANSCRIBE_SOURCES_SEPARATELY", True)
ENABLE_CHUNKED_TRANSCRIPTION = env_bool("ENABLE_CHUNKED_TRANSCRIPTION", True)
CHUNK_LENGTH_SECONDS = env_float("CHUNK_LENGTH_SECONDS", 30.0)
CHUNK_OVERLAP_SECONDS = env_float("CHUNK_OVERLAP_SECONDS", 3.0)
ENABLE_DIARIZATION = env_bool("ENABLE_DIARIZATION", False)
DIARIZATION_BACKEND = os.getenv("DIARIZATION_BACKEND", "pyannote").strip().lower()
HF_TOKEN = os.getenv("HF_TOKEN", "")
PYANNOTE_MODEL = os.getenv("PYANNOTE_MODEL", "pyannote/speaker-diarization-3.1")
MIN_SPEAKERS = os.getenv("MIN_SPEAKERS")
MAX_SPEAKERS = os.getenv("MAX_SPEAKERS")
DIARIZATION_TRANSCRIBE_TURNS = env_bool("DIARIZATION_TRANSCRIBE_TURNS", ENABLE_DIARIZATION)
WHISPERX_MODEL = os.getenv("WHISPERX_MODEL", "large-v3")
WHISPERX_DEVICE = os.getenv("WHISPERX_DEVICE", DEFAULT_DEVICE)
WHISPERX_COMPUTE_TYPE = os.getenv("WHISPERX_COMPUTE_TYPE", DEFAULT_COMPUTE_TYPE)
WHISPERX_BATCH_SIZE = env_int("WHISPERX_BATCH_SIZE", 8)
ENHANCE_UPLOADED_AUDIO = env_bool("ENHANCE_UPLOADED_AUDIO", False)
SYSTEM_AUDIO_SAMPLE_RATE = 44100
RECORDING_CHUNK_SECONDS = 0.25
TARGET_SOURCE_RMS = 0.08
TARGET_FINAL_RMS = 0.10
SUPPORTED_EXTENSIONS = ["wav", "mp3", "m4a", "mp4", "webm", "ogg", "flac", "aac"]
SUPPORTED_BACKENDS = {"faster_whisper", "whisperx", "auto", "nemo"}
WHISPERX_AVAILABLE = module_available("whisperx")
PYANNOTE_AVAILABLE = module_available("pyannote.audio")
NEMO_AVAILABLE = module_available("nemo") or module_available("nemo_toolkit")


# -----------------------------
# UI color constants
# Used by inject_css() to keep the Streamlit interface consistent.
# -----------------------------
BG_APP = "#F5F7FB"
BG_CARD = "#FFFFFF"
BG_NOTE = "#F8FAFD"
TEXT_MUTED = "#4B5563"
BORDER = "#D6DCE8"
TITLE = "#111827"
BUTTON_BG = "#E8ECEF"
BUTTON_BORDER = "#A7AFBC"


# -----------------------------
# Local asset paths
# These files are used for app icon, help image, and downloadable PDF.
# -----------------------------
BASE_DIR = Path(__file__).resolve().parent
ASSET_DIR = BASE_DIR / "assets"
ICON_PATH = ASSET_DIR / "app_icon.png"
HELP_IMAGE_PATH = ASSET_DIR / "Help_picture.png"
HELP_PDF_PATH = ASSET_DIR / "help_audio_to_text.pdf"



# -----------------------------
# Data container for transcription output
# Keeps transcript text, language, duration, and time-stamped segments together.
# -----------------------------
@dataclass
class TranscriptResult:
    text: str
    language: str | None
    duration: float | None
    segments: list[dict[str, Any]]


class WhisperTranscriber:
    def __init__(self, model_size: str = DEFAULT_MODEL_SIZE) -> None:
        self.model_size = model_size
        self._model: Any | None = None

    def get_model(self) -> Any:
        if WhisperModel is None:
            raise RuntimeError(
                "faster-whisper is not installed. Install requirements_streamlit.txt before running the app."
            )
        # Lazy-load model on first use to avoid slow startup.
        if self._model is None:
            self._model = WhisperModel(
                self.model_size,
                device=DEFAULT_DEVICE,
                compute_type=DEFAULT_COMPUTE_TYPE,
            )
        return self._model

    def transcribe(self, audio_path: str) -> TranscriptResult:
        model = self.get_model()
        transcribe_kwargs = {
            "language": WHISPER_LANGUAGE,
            "vad_filter": WHISPER_VAD_FILTER,
            "beam_size": WHISPER_BEAM_SIZE,
            "best_of": WHISPER_BEST_OF,
            "temperature": WHISPER_TEMPERATURE,
            "no_speech_threshold": WHISPER_NO_SPEECH_THRESHOLD,
            "log_prob_threshold": WHISPER_LOG_PROB_THRESHOLD,
            "condition_on_previous_text": WHISPER_CONDITION_ON_PREVIOUS_TEXT,
        }
        # Lower no-speech filtering and extra VAD padding reduce the chance that
        # quieter or overlapping speakers are discarded in mixed recordings.
        # Tune VAD to keep quiet or overlapping speech instead of dropping it.
        if WHISPER_VAD_FILTER:
            transcribe_kwargs["vad_parameters"] = {
                "threshold": 0.35,
                "min_silence_duration_ms": 700,
                "speech_pad_ms": 400,
            }

        # Run ASR and collect generator output into a list of segment dictionaries.
        segments, info = model.transcribe(audio_path, **transcribe_kwargs)

        collected_segments: list[dict[str, Any]] = []
        full_text_parts: list[str] = []
        for seg in segments:
            seg_text = (seg.text or "").strip()
            collected_segments.append(
                {
                    "start": round(seg.start, 2),
                    "end": round(seg.end, 2),
                    "text": seg_text,
                }
            )
            if seg_text:
                full_text_parts.append(seg_text)

        transcript = " ".join(full_text_parts).strip()
        return TranscriptResult(
            text=transcript,
            language=getattr(info, "language", None),
            duration=getattr(info, "duration", None),
            segments=collected_segments,
        )



# Cache the transcriber so the Whisper model is loaded only once per Streamlit session.
@st.cache_resource(show_spinner=False)
def get_transcriber(model_size: str) -> WhisperTranscriber:
    return WhisperTranscriber(model_size=model_size)



# Initialize all Streamlit session-state keys used by the app.
# Debug tip: check these keys if UI values are not updating as expected.
def init_state() -> None:
    defaults = {
        "status": "Ready",
        "current_audio_bytes": None,
        "current_audio_name": "",
        "current_audio_source": "",
        "current_audio_format": "",
        "recorded_audio_bytes": None,
        "last_transcript": "",
        "last_segments": [],
        "last_language": None,
        "last_duration": None,
        "last_timestamp": "",
        "play_recording": False,
        "uploader_nonce": 0,
        "recorder_nonce": 0,
        "last_recording_id": None,
        "is_recording": False,
        "recording_session": None,
        "recording_started_at": None,
        "recording_messages": [],
        "last_audio_enhancement_status": "Audio enhancement has not run yet.",
        "last_transcription_mode": "Multi-speaker optimized decoding enabled.",
        "diarization_status": "Speaker diarization disabled. Set ENABLE_DIARIZATION=1 to enable optional diarization.",
        "recorded_source_audio_bytes": {},
        "last_backend_status": "Backend ready: faster_whisper.",
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def read_bytes(path: Path) -> bytes:
    return path.read_bytes()


def to_base64(path: Path) -> str:
    return base64.b64encode(read_bytes(path)).decode("utf-8")


def clickable_help_html(image_path: Path, pdf_path: Path) -> str:
    img_b64 = to_base64(image_path)
    pdf_b64 = to_base64(pdf_path)
    return f'''
    <div class="help-wrap">
      <a href="data:application/pdf;base64,{pdf_b64}" download="help_audio_to_text.pdf" title="Download help PDF">
        <img src="data:image/png;base64,{img_b64}" alt="Help" class="help-image" />
      </a>
    </div>
    '''


def build_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Audio Transcript", level=1)
    meta = document.add_paragraph()
    meta.add_run("Source: ").bold = True
    meta.add_run(st.session_state.current_audio_source or "audio")
    meta.add_run("\nFilename: ").bold = True
    meta.add_run(st.session_state.current_audio_name or "N/A")
    meta.add_run("\nGenerated: ").bold = True
    meta.add_run(st.session_state.last_timestamp or "N/A")
    meta.add_run("\nLanguage: ").bold = True
    meta.add_run(st.session_state.last_language or "N/A")
    meta.add_run("\nDuration (s): ").bold = True
    duration_text = (
        str(st.session_state.last_duration)
        if st.session_state.last_duration is not None
        else "N/A"
    )
    meta.add_run(duration_text)

    document.add_heading("Transcript", level=2)
    document.add_paragraph(st.session_state.last_transcript)

    if st.session_state.last_segments:
        document.add_heading("Time-stamped Segments", level=2)
        for seg in st.session_state.last_segments:
            document.add_paragraph(format_segment_line(seg))

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_stem() -> str:
    name = st.session_state.current_audio_name.strip()
    if not name:
        return "audio"
    return Path(name).stem or "audio"


def set_current_audio(audio_bytes: bytes, name: str, source: str, audio_format: str) -> None:
    st.session_state.current_audio_bytes = audio_bytes
    st.session_state.current_audio_name = name
    st.session_state.current_audio_source = source
    st.session_state.current_audio_format = audio_format
    is_recording = source != "uploaded audio"
    st.session_state.status = (
        "Recording captured. Click Transcribe to continue."
        if is_recording
        else "Audio file selected. Click Transcribe to continue."
    )
    if not is_recording:
        st.session_state.recorded_audio_bytes = None
        st.session_state.recorded_source_audio_bytes = {}
        st.session_state.play_recording = False



# Validate packages needed for desktop microphone + system-audio recording.
def ensure_desktop_audio_dependencies() -> None:
    if sc is None or np is None:
        raise RuntimeError(
            "Combined recording requires soundcard and numpy. Install requirements_streamlit.txt again."
        )



# Try to find a loopback/monitor device for capturing computer audio.
# Debug tip: if system audio is missing, inspect soundcard device names here.
def default_speaker_loopback() -> Any | None:
    ensure_desktop_audio_dependencies()
    try:
        speaker = sc.default_speaker()
        if speaker is not None:
            return sc.get_microphone(speaker.name, include_loopback=True)
    except Exception:
        pass

    try:
        loopbacks = sc.all_microphones(include_loopback=True)
    except Exception:
        return None

    for microphone in loopbacks:
        name = getattr(microphone, "name", "").lower()
        if any(marker in name for marker in ("loopback", "monitor", "stereo mix", "output")):
            return microphone
    return loopbacks[0] if loopbacks else None



# Discover microphone and computer-audio devices for combined recording.
def discover_recording_devices() -> tuple[list[tuple[str, Any]], list[str]]:
    ensure_desktop_audio_dependencies()
    devices: list[tuple[str, Any]] = []
    messages: list[str] = []

    try:
        microphone = sc.default_microphone()
        if microphone is not None:
            devices.append(("microphone", microphone))
        else:
            messages.append("No default microphone was found.")
    except Exception as exc:
        messages.append(f"Microphone is unavailable: {exc}")

    speaker_loopback = default_speaker_loopback()
    if speaker_loopback is not None:
        devices.append(("computer audio", speaker_loopback))
    else:
        messages.append(
            "Computer audio loopback was not found. On Windows, enable WASAPI loopback or Stereo Mix. "
            "On Linux, check that PulseAudio/PipeWire monitor sources are available."
        )

    if not devices:
        raise RuntimeError("No microphone or computer audio device was found on this machine.")
    return devices, messages



# -----------------------------
# Audio normalization and enhancement helpers
# These functions clean and prepare audio before transcription.
# -----------------------------
def to_float32_audio(audio: Any) -> np.ndarray:
    data = np.asarray(audio, dtype=np.float32)
    if data.ndim == 1:
        data = data[:, None]
    if data.size == 0:
        return np.zeros((0, 1), dtype=np.float32)
    return np.nan_to_num(data, copy=False)


def normalize_audio(audio: Any) -> np.ndarray:
    return to_float32_audio(audio)


def remove_dc_offset(audio: np.ndarray) -> np.ndarray:
    audio = to_float32_audio(audio).copy()
    if audio.size:
        audio -= np.mean(audio, axis=0, keepdims=True)
    return audio


def safe_peak_limit(audio: np.ndarray, peak: float = 0.95) -> np.ndarray:
    audio = to_float32_audio(audio)
    current_peak = float(np.max(np.abs(audio))) if audio.size else 0.0
    if current_peak > peak and current_peak > 0:
        audio = audio * (peak / current_peak)
    return np.clip(audio, -peak, peak)


def rms_level(audio: np.ndarray) -> float:
    audio = to_float32_audio(audio)
    return float(np.sqrt(np.mean(audio**2))) if audio.size else 0.0


def normalize_rms(audio: np.ndarray, target_rms: float = TARGET_SOURCE_RMS) -> np.ndarray:
    audio = to_float32_audio(audio)
    current_rms = rms_level(audio)
    if current_rms <= 1e-6:
        return audio
    gain = min(8.0, max(0.15, target_rms / current_rms))
    return safe_peak_limit(audio * gain)


def highpass_filter(audio: np.ndarray, sample_rate: int = SYSTEM_AUDIO_SAMPLE_RATE, cutoff_hz: float = 80.0) -> np.ndarray:
    audio = remove_dc_offset(audio)
    if not audio.size:
        return audio

    # Prefer scipy high-pass filter for longer audio; otherwise use manual fallback.
    if signal is not None and audio.shape[0] > sample_rate // 10:
        try:
            sos = signal.butter(4, cutoff_hz, btype="highpass", fs=sample_rate, output="sos")
            return signal.sosfiltfilt(sos, audio, axis=0).astype(np.float32)
        except Exception:
            pass

    rc = 1.0 / (2.0 * np.pi * cutoff_hz)
    dt = 1.0 / sample_rate
    alpha = rc / (rc + dt)
    filtered = np.zeros_like(audio)
    filtered[0] = audio[0]
    for idx in range(1, audio.shape[0]):
        filtered[idx] = alpha * (filtered[idx - 1] + audio[idx] - audio[idx - 1])
    return filtered


def estimate_noise_profile(audio: np.ndarray, sample_rate: int = SYSTEM_AUDIO_SAMPLE_RATE) -> np.ndarray:
    audio = to_float32_audio(audio)
    frame_size = max(1, int(0.03 * sample_rate))
    mono = np.mean(np.abs(audio), axis=1)
    usable = mono[: (len(mono) // frame_size) * frame_size]
    if not usable.size:
        return np.array([0.0], dtype=np.float32)
    frame_rms = np.sqrt(np.mean(usable.reshape(-1, frame_size) ** 2, axis=1))
    return frame_rms.astype(np.float32)



# Lightweight noise reduction used when scipy spectral cleanup is unavailable.
def fallback_noise_reduction(audio: np.ndarray, sample_rate: int = SYSTEM_AUDIO_SAMPLE_RATE) -> np.ndarray:
    filtered = highpass_filter(audio, sample_rate=sample_rate)
    frame_rms = estimate_noise_profile(filtered, sample_rate=sample_rate)
    noise_floor = float(np.percentile(frame_rms, 20)) if frame_rms.size else 0.0
    if noise_floor > 0:
        frame_size = max(1, int(0.02 * sample_rate))
        mono = np.mean(np.abs(filtered), axis=1)
        usable = mono[: (len(mono) // frame_size) * frame_size]
        if usable.size:
            envelope = np.sqrt(np.mean(usable.reshape(-1, frame_size) ** 2, axis=1))
            envelope = np.repeat(envelope, frame_size)
            if envelope.size < len(mono):
                envelope = np.pad(envelope, (0, len(mono) - envelope.size), mode="edge")
            gate = np.clip((envelope / (noise_floor * 2.2))[:, None], 0.20, 1.0)
            filtered *= gate
    return normalize_rms(safe_peak_limit(filtered), TARGET_FINAL_RMS)



# Advanced spectral denoising using scipy STFT/ISTFT.
# Debug tip: failures here automatically fall back to fallback_noise_reduction().
def spectral_noise_reduction(audio: np.ndarray, sample_rate: int = SYSTEM_AUDIO_SAMPLE_RATE) -> np.ndarray:
    if signal is None:
        raise RuntimeError("scipy is unavailable")

    audio = highpass_filter(audio, sample_rate=sample_rate)
    enhanced_channels: list[np.ndarray] = []
    nperseg = min(2048, max(256, audio.shape[0] // 4))
    noverlap = nperseg // 2

    for channel_idx in range(audio.shape[1]):
        channel = audio[:, channel_idx]
        freqs, times, spectrum = signal.stft(
            channel,
            fs=sample_rate,
            nperseg=nperseg,
            noverlap=noverlap,
            boundary="zeros",
        )
        magnitude = np.abs(spectrum)
        phase = np.exp(1j * np.angle(spectrum))
        frame_energy = np.mean(magnitude, axis=0)
        if frame_energy.size:
            quiet_limit = np.percentile(frame_energy, 25)
            quiet_frames = magnitude[:, frame_energy <= quiet_limit]
            noise_mag = np.median(quiet_frames, axis=1, keepdims=True) if quiet_frames.size else np.median(magnitude, axis=1, keepdims=True)
        else:
            noise_mag = np.zeros((magnitude.shape[0], 1), dtype=np.float32)

        reduction_strength = 1.35
        floor = 0.18
        clean_mag = np.maximum(magnitude - reduction_strength * noise_mag, floor * magnitude)
        cleaned_spectrum = clean_mag * phase
        _, cleaned = signal.istft(cleaned_spectrum, fs=sample_rate, nperseg=nperseg, noverlap=noverlap)
        if cleaned.shape[0] < audio.shape[0]:
            cleaned = np.pad(cleaned, (0, audio.shape[0] - cleaned.shape[0]))
        enhanced_channels.append(cleaned[: audio.shape[0]].astype(np.float32))

    enhanced = np.stack(enhanced_channels, axis=1)
    return normalize_rms(safe_peak_limit(enhanced), TARGET_FINAL_RMS)



# Main audio enhancement wrapper used before transcription.
# Returns both enhanced audio and human-readable status/debug messages.
def enhance_audio_for_transcription(
    audio: np.ndarray,
    sample_rate: int = SYSTEM_AUDIO_SAMPLE_RATE,
    *,
    allow_advanced: bool = True,
) -> tuple[np.ndarray, list[str]]:
    audio = to_float32_audio(audio)
    messages: list[str] = []
    if not audio.size:
        return audio, ["No audio was available for enhancement."]

    try:
        # Try advanced spectral denoising first when scipy is available.
        if allow_advanced and signal is not None:
            enhanced = spectral_noise_reduction(audio, sample_rate=sample_rate)
            messages.append("Advanced noise suppression applied.")
        else:
            raise RuntimeError("scipy is unavailable")
    except Exception as exc:
        # Safe fallback keeps app usable even when scipy/STFT cleanup fails.
        enhanced = fallback_noise_reduction(audio, sample_rate=sample_rate)
        messages.append(f"Advanced noise suppression unavailable; using fallback cleanup. Reason: {exc}")

    return safe_peak_limit(enhanced), messages


def reduce_recording_noise(audio: np.ndarray, sample_rate: int = SYSTEM_AUDIO_SAMPLE_RATE) -> np.ndarray:
    enhanced, _ = enhance_audio_for_transcription(audio, sample_rate=sample_rate)
    return enhanced


def prepare_source_track(track: np.ndarray, sample_rate: int = SYSTEM_AUDIO_SAMPLE_RATE) -> np.ndarray:
    cleaned = fallback_noise_reduction(track, sample_rate=sample_rate)
    return normalize_rms(cleaned, TARGET_SOURCE_RMS)


def mix_audio_tracks(tracks: list[np.ndarray]) -> np.ndarray:
    # Ignore empty/silent tracks before mixing.
    usable_tracks = [track for track in tracks if track.size and np.any(np.abs(track) > 1e-5)]
    if not usable_tracks:
        raise RuntimeError("No audio was captured from the microphone or computer output.")

    prepared_tracks = [prepare_source_track(track) for track in usable_tracks]
    max_frames = max(track.shape[0] for track in prepared_tracks)
    max_channels = max(track.shape[1] for track in prepared_tracks)
    aligned_tracks: list[np.ndarray] = []

    for track in prepared_tracks:
        if track.shape[1] == 1 and max_channels > 1:
            track = np.repeat(track, max_channels, axis=1)
        elif track.shape[1] > max_channels:
            track = track[:, :max_channels]
        aligned = np.zeros((max_frames, max_channels), dtype=np.float32)
        aligned[: track.shape[0], : track.shape[1]] = track
        aligned_tracks.append(aligned)

    mixed = np.mean(aligned_tracks, axis=0).astype(np.float32)
    return normalize_rms(safe_peak_limit(mixed), TARGET_FINAL_RMS)


def wav_bytes_from_float_audio(audio: np.ndarray, sample_rate: int = SYSTEM_AUDIO_SAMPLE_RATE) -> bytes:
    audio = safe_peak_limit(to_float32_audio(audio))
    pcm = (np.clip(audio, -1.0, 1.0) * 32767).astype(np.int16)
    buffer = io.BytesIO()
    # Write standard 16-bit PCM WAV bytes for playback/transcription.
    with wave.open(buffer, "wb") as wav_file:
        wav_file.setnchannels(pcm.shape[1])
        wav_file.setsampwidth(2)
        wav_file.setframerate(sample_rate)
        wav_file.writeframes(pcm.tobytes())
    return buffer.getvalue()


def wav_bytes_to_float_audio(audio_bytes: bytes) -> tuple[np.ndarray, int]:
    with wave.open(io.BytesIO(audio_bytes), "rb") as wav_file:
        channels = wav_file.getnchannels()
        sample_rate = wav_file.getframerate()
        sample_width = wav_file.getsampwidth()
        frames = wav_file.readframes(wav_file.getnframes())

    if sample_width != 2:
        raise RuntimeError("Only 16-bit WAV enhancement is supported for uploaded audio.")
    pcm = np.frombuffer(frames, dtype=np.int16).astype(np.float32) / 32767.0
    return pcm.reshape(-1, channels), sample_rate



# -----------------------------
# Desktop recording workflow
# Captures microphone/system audio in background threads.
# -----------------------------
def recorder_worker(source_name: str, device: Any, stop_event: threading.Event, state: dict[str, Any]) -> None:
    frames_per_chunk = max(1, int(SYSTEM_AUDIO_SAMPLE_RATE * RECORDING_CHUNK_SECONDS))
    try:
        with device.recorder(samplerate=SYSTEM_AUDIO_SAMPLE_RATE) as recorder:
            while not stop_event.is_set():
                chunk = normalize_audio(recorder.record(numframes=frames_per_chunk))
                with state["lock"]:
                    state["chunks"].setdefault(source_name, []).append(chunk)
    except Exception as exc:
        with state["lock"]:
            state["errors"].append(f"{source_name} capture failed: {exc}")



# Start background audio capture threads and store them in session state.
def start_recording_session() -> None:
    if st.session_state.is_recording:
        return

    devices, messages = discover_recording_devices()
    stop_event = threading.Event()
    state: dict[str, Any] = {
        "stop_event": stop_event,
        "threads": [],
        "chunks": {},
        "errors": [],
        "messages": messages,
        "started_at": time.time(),
        "lock": threading.Lock(),
    }

    for source_name, device in devices:
        thread = threading.Thread(
            target=recorder_worker,
            args=(source_name, device, stop_event, state),
            daemon=True,
        )
        thread.start()
        state["threads"].append(thread)

    st.session_state.recording_session = state
    st.session_state.is_recording = True
    st.session_state.recording_started_at = state["started_at"]
    st.session_state.recording_messages = messages
    st.session_state.status = "Recording microphone and computer audio. Click Stop Recording when finished."



# Stop recording, enhance captured tracks, mix sources, and return WAV bytes.
def stop_recording_session() -> tuple[bytes, str, list[str]]:
    state = st.session_state.recording_session
    if not state:
        raise RuntimeError("No active recording session was found.")

    # Signal all recorder threads to stop, then collect captured chunks.
    state["stop_event"].set()
    for thread in state["threads"]:
        thread.join(timeout=2.0)

    with state["lock"]:
        chunks_by_source = dict(state["chunks"])
        messages = list(state["messages"]) + list(state["errors"])

    st.session_state.recording_session = None
    st.session_state.is_recording = False
    st.session_state.recording_started_at = None

    tracks: list[np.ndarray] = []
    recorded_sources: list[str] = []
    source_audio_bytes: dict[str, bytes] = {}
    source_status: list[str] = []
    # Enhance each individual source for optional source-aware transcription.
    for source_name, chunks in chunks_by_source.items():
        if not chunks:
            continue
        track = normalize_audio(np.concatenate(chunks, axis=0))
        if track.size and np.any(np.abs(track) > 1e-5):
            enhanced_track, track_messages = enhance_audio_for_transcription(track)
            source_audio_bytes[source_name] = wav_bytes_from_float_audio(enhanced_track)
            source_status.extend(f"{source_name}: {message}" for message in track_messages)
            tracks.append(track)
            recorded_sources.append(source_name)

    # Mix microphone and computer audio into one main track.
    mixed_audio = mix_audio_tracks(tracks)
    cleaned_audio, enhancement_messages = enhance_audio_for_transcription(mixed_audio)
    messages.extend(source_status)
    messages.extend(enhancement_messages)
    st.session_state.recorded_source_audio_bytes = source_audio_bytes
    st.session_state.last_audio_enhancement_status = " ".join(enhancement_messages)
    return wav_bytes_from_float_audio(cleaned_audio), " + ".join(recorded_sources) or "recorded audio", messages


def discard_recording_session() -> None:
    state = st.session_state.get("recording_session")
    if state:
        state["stop_event"].set()
        for thread in state.get("threads", []):
            thread.join(timeout=1.0)
    st.session_state.recording_session = None
    st.session_state.is_recording = False
    st.session_state.recording_started_at = None



# Reset UI, audio input, recording state, and transcription output.
def clear_all() -> None:
    discard_recording_session()
    st.session_state.current_audio_bytes = None
    st.session_state.current_audio_name = ""
    st.session_state.current_audio_source = ""
    st.session_state.current_audio_format = ""
    st.session_state.recorded_audio_bytes = None
    st.session_state.recorded_source_audio_bytes = {}
    st.session_state.last_transcript = ""
    st.session_state.last_segments = []
    st.session_state.last_language = None
    st.session_state.last_duration = None
    st.session_state.last_timestamp = ""
    st.session_state.play_recording = False
    st.session_state.last_recording_id = None
    st.session_state.uploader_nonce += 1
    st.session_state.recorder_nonce += 1
    st.session_state.status = "Input and output cleared."
    st.rerun()




# Save selected/recorded audio bytes to a temporary file for ASR backends.
def save_current_audio_to_tempfile() -> str:
    audio_bytes = st.session_state.current_audio_bytes
    if not audio_bytes:
        raise RuntimeError("No audio is currently selected.")

    suffix = ".wav"
    name = st.session_state.current_audio_name or "audio.wav"
    ext = Path(name).suffix.lower()
    if ext in {f".{x}" for x in SUPPORTED_EXTENSIONS}:
        suffix = ext

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(audio_bytes)
        return tmp.name



# -----------------------------
# Transcript segment formatting and deduplication helpers
# Used to display clean, chronological transcript output.
# -----------------------------
def format_segment_label(seg: dict[str, Any]) -> str:
    labels = [str(seg.get(key, "")).strip() for key in ("speaker", "source") if seg.get(key)]
    return " / ".join(labels)


def format_segment_line(seg: dict[str, Any]) -> str:
    label = format_segment_label(seg)
    prefix = f"[{seg.get('start', 0):.2f}s -> {seg.get('end', 0):.2f}s]"
    if label:
        prefix += f" {label}:"
    return f"{prefix} {seg.get('text', '')}".strip()


def transcript_text_from_segments(segments: list[dict[str, Any]]) -> str:
    return "\n".join(format_segment_line(seg) for seg in segments if seg.get("text"))


def normalize_text_for_dedup(text: str) -> str:
    return " ".join("".join(ch.lower() if ch.isalnum() or ch.isspace() else " " for ch in text).split())


def segments_overlap(a: dict[str, Any], b: dict[str, Any], tolerance: float = 0.35) -> bool:
    return min(float(a.get("end", 0)), float(b.get("end", 0))) + tolerance >= max(float(a.get("start", 0)), float(b.get("start", 0)))



# Detect repeated text caused by source-aware or overlapping chunk transcription.
def is_duplicate_segment(a: dict[str, Any], b: dict[str, Any]) -> bool:
    a_text = normalize_text_for_dedup(str(a.get("text", "")))
    b_text = normalize_text_for_dedup(str(b.get("text", "")))
    if not a_text or not b_text:
        return False
    if not segments_overlap(a, b):
        return False
    similarity = difflib.SequenceMatcher(None, a_text, b_text).ratio()
    same_source = a.get("source") == b.get("source")
    short_repeat = len(a_text.split()) <= 4 and a_text == b_text
    return similarity >= (0.88 if same_source else 0.94) or short_repeat



# Sort transcript segments and remove near-duplicate overlapping segments.
def merge_segments_chronologically(segments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    # Sort by timestamp before removing near-duplicates.
    sorted_segments = sorted(
        (seg for seg in segments if str(seg.get("text", "")).strip()),
        key=lambda seg: (float(seg.get("start", 0)), float(seg.get("end", 0))),
    )
    merged: list[dict[str, Any]] = []
    for seg in sorted_segments:
        candidate = dict(seg)
        candidate["start"] = round(float(candidate.get("start", 0)), 2)
        candidate["end"] = round(float(candidate.get("end", candidate["start"])), 2)
        duplicate_index = next((idx for idx, existing in enumerate(merged) if is_duplicate_segment(candidate, existing)), None)
        if duplicate_index is None:
            merged.append(candidate)
            continue
        existing = merged[duplicate_index]
        if len(str(candidate.get("text", ""))) > len(str(existing.get("text", ""))):
            merged[duplicate_index] = candidate
    return merged


def audio_duration_from_wav_bytes(audio_bytes: bytes) -> float | None:
    try:
        with wave.open(io.BytesIO(audio_bytes), "rb") as wav_file:
            return wav_file.getnframes() / float(wav_file.getframerate())
    except Exception:
        return None


def write_temp_wav_segment(audio: np.ndarray, sample_rate: int, start: float, end: float) -> str:
    start_frame = max(0, int(start * sample_rate))
    end_frame = min(audio.shape[0], int(end * sample_rate))
    segment_audio = audio[start_frame:end_frame]
    if not segment_audio.size:
        segment_audio = np.zeros((1, audio.shape[1]), dtype=np.float32)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
        tmp.write(wav_bytes_from_float_audio(segment_audio, sample_rate=sample_rate))
        return tmp.name


def offset_segments(segments: list[dict[str, Any]], offset: float, *, source: str, backend: str, speaker: str | None = None) -> list[dict[str, Any]]:
    output: list[dict[str, Any]] = []
    # Assign each ASR segment to the speaker turn with the largest time overlap.
    for seg in segments:
        item = dict(seg)
        item["start"] = round(float(item.get("start", 0)) + offset, 2)
        item["end"] = round(float(item.get("end", 0)) + offset, 2)
        item["source"] = source
        item["backend"] = backend
        if speaker and not item.get("speaker"):
            item["speaker"] = speaker
        output.append(item)
    return output



# -----------------------------
# ASR transcription helpers
# Handles chunking, source labels, optional WhisperX/NeMo, and diarization.
# -----------------------------
# Transcribe long WAV files in overlapping chunks to reduce missed speech.
def transcribe_audio_chunked(audio_path: str, *, source: str, backend: str = "faster_whisper") -> TranscriptResult:
    try:
        audio = read_wav_file_to_float(audio_path)
    except Exception:
        result = get_transcriber(DEFAULT_MODEL_SIZE).transcribe(audio_path)
        result.segments = offset_segments(result.segments, 0.0, source=source, backend=backend)
        result.text = transcript_text_from_segments(result.segments)
        return result

    audio_data, sample_rate = audio
    duration = audio_data.shape[0] / float(sample_rate)
    if not ENABLE_CHUNKED_TRANSCRIPTION or duration <= CHUNK_LENGTH_SECONDS + CHUNK_OVERLAP_SECONDS:
        result = get_transcriber(DEFAULT_MODEL_SIZE).transcribe(audio_path)
        result.segments = offset_segments(result.segments, 0.0, source=source, backend=backend)
        result.text = transcript_text_from_segments(result.segments)
        return result

    all_segments: list[dict[str, Any]] = []
    step = max(1.0, CHUNK_LENGTH_SECONDS - CHUNK_OVERLAP_SECONDS)
    start = 0.0
    temp_paths: list[str] = []
    try:
        # Create overlapping chunks so speech near boundaries is not lost.
        while start < duration:
            end = min(duration, start + CHUNK_LENGTH_SECONDS)
            chunk_path = write_temp_wav_segment(audio_data, sample_rate, start, end)
            temp_paths.append(chunk_path)
            chunk_result = get_transcriber(DEFAULT_MODEL_SIZE).transcribe(chunk_path)
            all_segments.extend(offset_segments(chunk_result.segments, start, source=source, backend=backend))
            if end >= duration:
                break
            start += step
    finally:
        for temp_path in temp_paths:
            try:
                Path(temp_path).unlink()
            except Exception:
                pass

    # Merge overlap-created duplicates after chunk transcription.
    merged = merge_segments_chronologically(all_segments)
    return TranscriptResult(transcript_text_from_segments(merged), WHISPER_LANGUAGE, duration, merged)


def read_wav_file_to_float(path: str) -> tuple[np.ndarray, int]:
    with open(path, "rb") as handle:
        return wav_bytes_to_float_audio(handle.read())


def write_bytes_to_temp_audio(audio_bytes: bytes, suffix: str = ".wav") -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(audio_bytes)
        return tmp.name


def asr_quality_message() -> str:
    if ASR_QUALITY_PRESET == "cpu_fast":
        return "CPU fast preset: prioritizes speed; small/base models are recommended."
    if ASR_QUALITY_PRESET == "high_accuracy":
        return "High accuracy preset: consider large-v3 or large-v3-turbo and GPU for best results."
    if ASR_QUALITY_PRESET == "multispeaker":
        return "Multispeaker preset: source-aware and chunked transcription are recommended; diarization can be enabled with ENABLE_DIARIZATION=1."
    return "Balanced preset: preserves safe default model/device behavior."


def model_warning_message() -> str:
    model = DEFAULT_MODEL_SIZE.lower()
    if DEFAULT_DEVICE == "cpu" and any(name in model for name in ("large", "medium")):
        return "Large Whisper models on CPU can be slow; use GPU or a smaller model if latency is important."
    return ""



# Choose ASR backend based on TRANSCRIPTION_BACKEND and installed packages.
def selected_backend() -> str:
    backend = TRANSCRIPTION_BACKEND if TRANSCRIPTION_BACKEND in SUPPORTED_BACKENDS else "faster_whisper"
    # Auto mode prefers WhisperX if installed; otherwise uses faster-whisper.
    if backend == "auto":
        return "whisperx" if WHISPERX_AVAILABLE else "faster_whisper"
    return backend



# Optional WhisperX backend; returns None if unavailable or failed.
def transcribe_with_whisperx(audio_path: str, *, source: str) -> TranscriptResult | None:
    if not WHISPERX_AVAILABLE:
        st.session_state.last_backend_status = "WhisperX is not installed; falling back to faster_whisper."
        return None
    try:
        import whisperx  # type: ignore

        model = whisperx.load_model(WHISPERX_MODEL, WHISPERX_DEVICE, compute_type=WHISPERX_COMPUTE_TYPE)
        result = model.transcribe(audio_path, batch_size=WHISPERX_BATCH_SIZE, language=WHISPER_LANGUAGE)
        segments = []
        for seg in result.get("segments", []):
            segments.append(
                {
                    "start": round(float(seg.get("start", 0)), 2),
                    "end": round(float(seg.get("end", 0)), 2),
                    "text": str(seg.get("text", "")).strip(),
                    "source": source,
                    "backend": "whisperx",
                }
            )
        st.session_state.last_backend_status = "WhisperX ASR backend used."
        return TranscriptResult(transcript_text_from_segments(segments), result.get("language"), None, segments)
    except Exception as exc:
        st.session_state.last_backend_status = f"WhisperX failed; falling back to faster_whisper. Reason: {exc}"
        return None



# Placeholder for future NVIDIA NeMo ASR integration.
def transcribe_with_nemo(audio_path: str, *, source: str) -> TranscriptResult | None:
    st.session_state.last_backend_status = "NeMo backend is prepared but not implemented; falling back to faster_whisper."
    return None



# Optional speaker diarization using pyannote.audio.
# Debug tip: requires ENABLE_DIARIZATION=1 and HF_TOKEN for pyannote models.
def diarization_turns(audio_path: str) -> list[dict[str, Any]]:
    # Diarization has several optional dependencies/settings, so fail softly.
    if not ENABLE_DIARIZATION:
        st.session_state.diarization_status = "Diarization disabled."
        return []
    if not PYANNOTE_AVAILABLE:
        st.session_state.diarization_status = "pyannote.audio is not installed; continuing without diarization."
        return []
    if DIARIZATION_BACKEND not in {"pyannote", "whisperx"}:
        st.session_state.diarization_status = f"Unsupported diarization backend '{DIARIZATION_BACKEND}'; continuing without diarization."
        return []
    if not HF_TOKEN and DIARIZATION_BACKEND == "pyannote":
        st.session_state.diarization_status = "HF_TOKEN is not set for pyannote diarization; continuing without diarization."
        return []

    try:
        from pyannote.audio import Pipeline  # type: ignore

        pipeline = Pipeline.from_pretrained(PYANNOTE_MODEL, use_auth_token=HF_TOKEN or None)
        kwargs: dict[str, int] = {}
        if MIN_SPEAKERS:
            kwargs["min_speakers"] = int(MIN_SPEAKERS)
        if MAX_SPEAKERS:
            kwargs["max_speakers"] = int(MAX_SPEAKERS)
        diarization = pipeline(audio_path, **kwargs)
        turns = []
        for turn, _, speaker in diarization.itertracks(yield_label=True):
            turns.append({"start": float(turn.start), "end": float(turn.end), "speaker": str(speaker)})
        st.session_state.diarization_status = f"Diarization enabled: {len(turns)} speaker turns found."
        return turns
    except Exception as exc:
        st.session_state.diarization_status = f"Diarization failed; continuing without diarization. Reason: {exc}"
        return []


def transcribe_speaker_turns(audio_path: str, turns: list[dict[str, Any]], *, source: str) -> list[dict[str, Any]]:
    if not turns or not DIARIZATION_TRANSCRIBE_TURNS:
        return []
    try:
        audio, sample_rate = read_wav_file_to_float(audio_path)
    except Exception:
        return []

    duration = audio.shape[0] / float(sample_rate)
    all_segments: list[dict[str, Any]] = []
    temp_paths: list[str] = []
    try:
        for turn in turns:
            start = max(0.0, float(turn["start"]) - 0.5)
            end = min(duration, float(turn["end"]) + 0.5)
            chunk_path = write_temp_wav_segment(audio, sample_rate, start, end)
            temp_paths.append(chunk_path)
            result = get_transcriber(DEFAULT_MODEL_SIZE).transcribe(chunk_path)
            all_segments.extend(
                offset_segments(result.segments, start, source=source, backend="faster_whisper", speaker=str(turn.get("speaker", "")) or None)
            )
    finally:
        for temp_path in temp_paths:
            try:
                Path(temp_path).unlink()
            except Exception:
                pass
    return all_segments



# Transcribe one audio source using selected backend with safe fallback.
def transcribe_single_source(audio_path: str, *, source: str) -> TranscriptResult:
    backend = selected_backend()
    if backend == "whisperx":
        result = transcribe_with_whisperx(audio_path, source=source)
        if result is not None:
            return result
    elif backend == "nemo":
        result = transcribe_with_nemo(audio_path, source=source)
        if result is not None:
            return result

    turns = diarization_turns(audio_path) if source == "mixed" else []
    if turns and DIARIZATION_TRANSCRIBE_TURNS:
        segments = transcribe_speaker_turns(audio_path, turns, source=source)
        if segments:
            merged = merge_segments_chronologically(segments)
            return TranscriptResult(transcript_text_from_segments(merged), WHISPER_LANGUAGE, None, merged)

    result = transcribe_audio_chunked(audio_path, source=source, backend="faster_whisper")
    if turns:
        result.segments = maybe_run_diarization(audio_path, result.segments)
        result.text = transcript_text_from_segments(result.segments)
    if "falling back" not in str(st.session_state.get("last_backend_status", "")).lower():
        st.session_state.last_backend_status = "faster_whisper backend used."
    return result



# Build list of audio sources to transcribe: mixed audio plus optional separate tracks.
def transcription_sources_for_current_audio(main_temp_path: str) -> list[tuple[str, str]]:
    sources = [("mixed", main_temp_path)]
    if not TRANSCRIBE_SOURCES_SEPARATELY:
        return sources
    source_audio = st.session_state.get("recorded_source_audio_bytes") or {}
    for source_name, audio_bytes in source_audio.items():
        label = "system" if source_name == "computer audio" else source_name
        try:
            sources.append((label, write_bytes_to_temp_audio(audio_bytes, ".wav")))
        except Exception:
            pass
    return sources



# Transcribe all selected sources and merge them into one chronological transcript.
def transcribe_current_audio_advanced(main_temp_path: str) -> TranscriptResult:
    source_paths = transcription_sources_for_current_audio(main_temp_path)
    temp_source_paths = [path for _, path in source_paths if path != main_temp_path]
    all_segments: list[dict[str, Any]] = []
    language = WHISPER_LANGUAGE
    duration = None
    try:
        for source, path in source_paths:
            result = transcribe_single_source(path, source=source)
            language = result.language or language
            duration = max(duration or 0, result.duration or 0) or duration
            all_segments.extend(result.segments)
    finally:
        for path in temp_source_paths:
            try:
                Path(path).unlink()
            except Exception:
                pass

    # Merge overlap-created duplicates after chunk transcription.
    merged = merge_segments_chronologically(all_segments)
    return TranscriptResult(transcript_text_from_segments(merged), language, duration, merged)



# Optionally enhance uploaded WAV audio before ASR.
def maybe_enhance_uploaded_audio_before_transcription() -> None:
    if not ENHANCE_UPLOADED_AUDIO:
        return
    if st.session_state.current_audio_source != "uploaded audio":
        return
    if (st.session_state.current_audio_format or "").lower() != "wav":
        st.session_state.last_audio_enhancement_status = (
            "Uploaded audio enhancement skipped because only WAV uploads can be enhanced in-app."
        )
        return

    try:
        audio, sample_rate = wav_bytes_to_float_audio(st.session_state.current_audio_bytes)
        enhanced, messages = enhance_audio_for_transcription(audio, sample_rate=sample_rate)
        st.session_state.current_audio_bytes = wav_bytes_from_float_audio(enhanced, sample_rate=sample_rate)
        st.session_state.current_audio_name = f"{export_stem()}_enhanced.wav"
        st.session_state.current_audio_format = "wav"
        st.session_state.last_audio_enhancement_status = " ".join(messages)
    except Exception as exc:
        st.session_state.last_audio_enhancement_status = f"Uploaded audio enhancement skipped: {exc}"


def maybe_run_diarization(audio_path: str, segments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    turns = diarization_turns(audio_path)
    if not turns:
        return segments

    labeled: list[dict[str, Any]] = []
    # Assign each ASR segment to the speaker turn with the largest time overlap.
    for seg in segments:
        item = dict(seg)
        best_turn = None
        best_overlap = 0.0
        for turn in turns:
            overlap = max(0.0, min(float(item.get("end", 0)), turn["end"]) - max(float(item.get("start", 0)), turn["start"]))
            if overlap > best_overlap:
                best_overlap = overlap
                best_turn = turn
        if best_turn and best_overlap > 0:
            item["speaker"] = best_turn["speaker"]
        labeled.append(item)
    return labeled



# Save transcription output into Streamlit session state for display/download.
def apply_transcript_result(result: TranscriptResult) -> None:
    st.session_state.last_segments = merge_segments_chronologically(result.segments)
    st.session_state.last_transcript = transcript_text_from_segments(st.session_state.last_segments)
    st.session_state.last_language = result.language
    st.session_state.last_duration = result.duration
    st.session_state.last_timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    st.session_state.status = "Transcription complete. Multi-speaker optimized decoding was used."



# Main transcription entry point called when the user clicks Transcribe.
def transcribe_current_audio() -> None:
    if not st.session_state.current_audio_bytes:
        st.session_state.status = "Please open an audio file or record audio first."
        return

    temp_path = None
    try:
        maybe_enhance_uploaded_audio_before_transcription()
        temp_path = save_current_audio_to_tempfile()
        st.session_state.status = (
            f"Transcribing with model: {DEFAULT_MODEL_SIZE}, device: {DEFAULT_DEVICE}, "
            f"compute type: {DEFAULT_COMPUTE_TYPE}. Multi-speaker optimized decoding enabled."
        )
        result = transcribe_current_audio_advanced(temp_path)
        apply_transcript_result(result)
    finally:
        if temp_path and Path(temp_path).exists():
            try:
                Path(temp_path).unlink()
            except Exception:
                pass


def summary_text() -> str:
    word_count = len(st.session_state.last_transcript.split()) if st.session_state.last_transcript else 0
    source = st.session_state.current_audio_source or "N/A"
    last_run = st.session_state.last_timestamp or "N/A"
    return f"Words: {word_count} | Segments: {len(st.session_state.last_segments)} | Source: {source} | Last run: {last_run}"


def selected_audio_label() -> str:
    if st.session_state.current_audio_name:
        return f"Selected: {st.session_state.current_audio_name}"
    return "No audio selected"



# -----------------------------
# Streamlit UI rendering helpers
# These functions define the layout, controls, status, and outputs.
# -----------------------------
def inject_css() -> None:
    st.markdown(
        f"""
        <style>
        .stApp {{
            background: {BG_APP};
        }}
        header[data-testid="stHeader"] {{
            display: none;
        }}
        div[data-testid="stToolbar"] {{
            display: none;
        }}
        #MainMenu {{
            visibility: hidden;
        }}
        .stAppViewContainer {{
            padding-top: 0;
        }}
        .main .block-container {{
            max-width: 1320px;
            padding-top: 0.55rem;
            padding-bottom: 1rem;
        }}
        .app-header {{
            display: flex;
            justify-content: space-between;
            gap: 1rem;
            align-items: flex-start;
            margin-bottom: 0.7rem;
        }}
        .brand-row {{
            display: flex;
            gap: 0.9rem;
            align-items: flex-start;
        }}
        .brand-icon {{
            width: 44px;
            height: 44px;
            object-fit: contain;
            margin-top: 4px;
        }}
        .brand-title {{
            font-size: 2.0rem;
            font-weight: 700;
            color: {TITLE};
            line-height: 1.1;
            margin: 0;
        }}
        .brand-subtitle {{
            color: {TEXT_MUTED};
            font-size: 0.98rem;
            margin-top: 0.35rem;
        }}
        .header-right {{
            display: flex;
            align-items: flex-start;
            gap: 0.9rem;
        }}
        .note-box {{
            background: {BG_NOTE};
            border: 1px solid {BORDER};
            padding: 0.65rem 0.9rem;
            border-radius: 0;
            font-size: 0.92rem;
            color: #1F2937;
            font-style: italic;
            text-align: center;
            min-width: 290px;
        }}
        .help-wrap {{
            text-align: center;
        }}
        .help-image {{
            width: 116px;
            height: auto;
            display: block;
        }}
        .card-title {{
            font-weight: 700;
            font-size: 1.03rem;
            color: {TITLE};
            margin-bottom: 0.65rem;
        }}
        div[data-testid="stVerticalBlockBorderWrapper"] {{
            background: {BG_CARD};
            border: 1px solid {BORDER};
            border-radius: 0;
            padding: 0.35rem 0.35rem 0.2rem 0.35rem;
            margin-bottom: 0.75rem;
        }}
        .section-label {{
            font-weight: 700;
            margin-top: 0.8rem;
            margin-bottom: 0.5rem;
            color: {TITLE};
        }}
        .muted-text {{
            color: {TEXT_MUTED};
            font-size: 0.94rem;
        }}
        .summary-line {{
            color: {TEXT_MUTED};
            font-size: 0.96rem;
            margin-top: 0.25rem;
            margin-bottom: 0.25rem;
        }}
        .subheading {{
            color: {TEXT_MUTED};
            font-size: 0.98rem;
            margin-top: 0.6rem;
            margin-bottom: 0.35rem;
        }}
        .stButton > button, .stDownloadButton > button {{
            width: 100%;
            background: {BUTTON_BG};
            color: #111827;
            border: 1px solid {BUTTON_BORDER};
            border-radius: 0;
            min-height: 2.6rem;
            font-size: 0.98rem;
            box-shadow: none;
        }}
        .stButton > button:hover, .stDownloadButton > button:hover {{
            border-color: #8A94A3;
            background: #EFF2F5;
            color: #111827;
        }}
        .stButton > button:disabled {{
            color: #8B93A1;
            background: #ECECEC;
        }}
        .transcript-box {{
            border: 1px solid {BORDER};
            background: #FFFFFF;
            min-height: 260px;
            max-height: 260px;
            overflow-y: auto;
            padding: 0.85rem;
            white-space: pre-wrap;
            color: {TITLE};
            font-size: 1rem;
        }}
        .status-box {{
            color: {TITLE};
            font-size: 0.98rem;
            margin-bottom: 0.5rem;
        }}
        .meta-box {{
            color: {TEXT_MUTED};
            font-size: 0.92rem;
        }}
        div[data-testid="stFileUploader"] section {{
            border: 1px solid {BUTTON_BORDER};
            border-radius: 0;
            background: #FAFBFC;
        }}
        div[data-testid="stFileUploader"] button {{
            border-radius: 0;
        }}
        div[data-testid="stAudioInput"] button {{
            width: 100%;
            border-radius: 0;
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_header() -> None:
    icon_b64 = to_base64(ICON_PATH)
    disclaimer = html.escape("Note: This code is created with the help of OpenAI\ncontact: arkabhowmik@yahoo.co.uk").replace(
        "\n", "<br>"
    )
    header_html = f'''
    <div class="app-header">
      <div class="brand-row">
        <img class="brand-icon" src="data:image/png;base64,{icon_b64}" alt="App icon">
        <div>
          <h1 class="brand-title">Audio to Text</h1>
          <div class="brand-subtitle">Desktop audio transcription using microphone recording or audio-file upload.</div>
        </div>
      </div>
      <div class="header-right">
        <div class="note-box">{disclaimer}</div>
        {clickable_help_html(HELP_IMAGE_PATH, HELP_PDF_PATH)}
      </div>
    </div>
    '''
    st.markdown(header_html, unsafe_allow_html=True)



# Render input controls: upload, recording, playback, actions, and status.
def render_left_panel() -> None:
    # Use bordered containers as card-style UI sections.
    with st.container(border=True):
        st.markdown('<div class="card-title">Input</div>', unsafe_allow_html=True)

        uploaded_file = st.file_uploader(
            "Open audio file",
            type=SUPPORTED_EXTENSIONS,
            key=f"uploader_{st.session_state.uploader_nonce}",
            label_visibility="visible",
            disabled=bool(st.session_state.is_recording),
        )
        # Store uploaded file bytes in session state for playback/transcription.
        if uploaded_file is not None:
            set_current_audio(
                audio_bytes=uploaded_file.getvalue(),
                name=uploaded_file.name,
                source="uploaded audio",
                audio_format=Path(uploaded_file.name).suffix.lower().lstrip(".") or "wav",
            )

        st.markdown(f'<div class="muted-text">{html.escape(selected_audio_label())}</div>', unsafe_allow_html=True)
        st.markdown('<div class="section-label">Recording</div>', unsafe_allow_html=True)

        combined_recorder_available = sc is not None and np is not None
        recorded_from_component = False

        if combined_recorder_available:
            if st.session_state.is_recording:
                started_at = st.session_state.recording_started_at or time.time()
                elapsed = int(time.time() - started_at)
                st.info(f"Recording active: {elapsed // 60:02d}:{elapsed % 60:02d}")
                for message in st.session_state.recording_messages:
                    st.warning(message)

                # Stop background recording and prepare enhanced WAV audio.
                if st.button("Stop Recording", key="combined_stop_recording_button"):
                    try:
                        with st.spinner("Recording stopped. Enhancing audio for transcription..."):
                            audio_bytes, recorded_sources, messages = stop_recording_session()
                        st.session_state.recorded_audio_bytes = audio_bytes
                        st.session_state.play_recording = False
                        set_current_audio(
                            audio_bytes=audio_bytes,
                            name=f"recorded_audio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.wav",
                            source="recorded audio",
                            audio_format="wav",
                        )
                        st.session_state.status = (
                            f"Recording captured from {recorded_sources}. Noise reduction applied. "
                            "Click Transcribe to continue."
                        )
                        for message in messages:
                            if "applied" in message.lower():
                                st.info(message)
                            else:
                                st.warning(message)
                    except Exception as exc:
                        discard_recording_session()
                        st.session_state.status = f"Recording failed: {exc}"
                        st.error(st.session_state.status)
            else:
                st.success("Recorder ready: microphone and computer audio will be captured together when available.")
                # Start desktop recording; rerun updates the UI timer/status.
                if st.button("Start Recording", key="combined_recording_button"):
                    try:
                        start_recording_session()
                        st.rerun()
                    except Exception as exc:
                        discard_recording_session()
                        st.session_state.status = f"Recording failed: {exc}"
                        st.error(st.session_state.status)
        elif mic_recorder is not None:
            st.warning(
                "Computer audio capture is unavailable because soundcard or numpy is not installed. "
                "Falling back to browser microphone recording."
            )
            mic_audio = mic_recorder(
                start_prompt="Start Recording",
                stop_prompt="Stop Recording",
                just_once=False,
                use_container_width=True,
                key=f"mic_{st.session_state.recorder_nonce}",
            )
            if mic_audio and isinstance(mic_audio, dict):
                recording_id = mic_audio.get("id")
                if recording_id != st.session_state.last_recording_id:
                    st.session_state.last_recording_id = recording_id
                    audio_bytes = mic_audio.get("bytes")
                    if audio_bytes:
                        st.session_state.recorded_audio_bytes = audio_bytes
                        st.session_state.play_recording = False
                        set_current_audio(
                            audio_bytes=audio_bytes,
                            name="recorded_audio.wav",
                            source="recorded audio",
                            audio_format="wav",
                        )
                        recorded_from_component = True
        else:
            st.warning(
                "Computer audio capture is unavailable because soundcard or numpy is not installed. "
                "Falling back to browser microphone recording."
            )
            audio_input = st.audio_input(
                "Start Recording",
                key=f"audio_input_{st.session_state.recorder_nonce}",
            )
            if audio_input is not None:
                audio_bytes = audio_input.getvalue()
                st.session_state.recorded_audio_bytes = audio_bytes
                st.session_state.play_recording = False
                set_current_audio(
                    audio_bytes=audio_bytes,
                    name="recorded_audio.wav",
                    source="recorded audio",
                    audio_format="wav",
                )
                recorded_from_component = True

        if recorded_from_component:
            st.session_state.status = "Recording captured. Click Transcribe to continue."

        play_clicked = st.button(
            "Play audio",
            disabled=not bool(st.session_state.current_audio_bytes),
            key="play_recording_button",
        )
        if play_clicked and st.session_state.current_audio_bytes:
            st.session_state.play_recording = True

        if st.session_state.play_recording and st.session_state.current_audio_bytes:
            audio_format = st.session_state.current_audio_format or "wav"
            st.audio(st.session_state.current_audio_bytes, format=f"audio/{audio_format}")

    # Use bordered containers as card-style UI sections.
    with st.container(border=True):
        st.markdown('<div class="card-title">Actions</div>', unsafe_allow_html=True)
        # Run ASR only when recording is not active.
        if st.button("Transcribe", key="transcribe_button", disabled=bool(st.session_state.is_recording)):
            try:
                with st.spinner("Loading model and transcribing audio..."):
                    transcribe_current_audio()
            except Exception as exc:
                st.session_state.status = f"Transcription failed: {exc}"
                st.error(st.session_state.status)
        if st.button("Clear output", key="clear_button"):
            clear_all()

    # Use bordered containers as card-style UI sections.
    with st.container(border=True):
        st.markdown('<div class="card-title">Status</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="status-box">{html.escape(st.session_state.status)}</div>', unsafe_allow_html=True)
        st.markdown(
            f'<div class="meta-box">Model: {html.escape(DEFAULT_MODEL_SIZE)} | Device: {html.escape(DEFAULT_DEVICE)} | Compute: {html.escape(DEFAULT_COMPUTE_TYPE)}</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            f'<div class="meta-box">Enhancement: {html.escape(st.session_state.last_audio_enhancement_status)}</div>',
            unsafe_allow_html=True,
        )
        backend_details = (
            f"Backend: {selected_backend()} (requested: {TRANSCRIPTION_BACKEND}) | "
            f"Preset: {ASR_QUALITY_PRESET} | Source-aware: {TRANSCRIBE_SOURCES_SEPARATELY} | "
            f"Chunked: {ENABLE_CHUNKED_TRANSCRIPTION} ({CHUNK_LENGTH_SECONDS:g}s/{CHUNK_OVERLAP_SECONDS:g}s overlap)"
        )
        optional_details = (
            f"Diarization: {ENABLE_DIARIZATION} | Turn ASR: {DIARIZATION_TRANSCRIBE_TURNS} | "
            f"WhisperX available: {WHISPERX_AVAILABLE} | pyannote available: {PYANNOTE_AVAILABLE}"
        )
        st.markdown(f'<div class="meta-box">{html.escape(backend_details)}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="meta-box">{html.escape(optional_details)}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="meta-box">{html.escape(asr_quality_message())}</div>', unsafe_allow_html=True)
        warning_message = model_warning_message()
        if warning_message:
            st.warning(warning_message)
        st.markdown(
            f'<div class="meta-box">Transcription: {html.escape(st.session_state.last_backend_status)} | {html.escape(st.session_state.diarization_status)}</div>',
            unsafe_allow_html=True,
        )



# Render transcript display, segment table, and download buttons.
def render_right_panel() -> None:
    export_col1, export_col2, _ = st.columns([1.1, 1.1, 5.8])
    txt_bytes = st.session_state.last_transcript.encode("utf-8") if st.session_state.last_transcript else b""
    with export_col1:
        st.download_button(
            "Download TXT",
            data=txt_bytes,
            file_name=f"{export_stem()}_transcript.txt",
            mime="text/plain",
            disabled=not bool(st.session_state.last_transcript),
            key="download_txt",
        )
    with export_col2:
        st.download_button(
            "Download DOCX",
            data=build_docx_bytes() if st.session_state.last_transcript else b"",
            file_name=f"{export_stem()}_transcript.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            disabled=not bool(st.session_state.last_transcript),
            key="download_docx",
        )

    # Use bordered containers as card-style UI sections.
    with st.container(border=True):
        st.markdown('<div class="card-title" style="margin-top:0.1rem;">Transcript</div>', unsafe_allow_html=True)
        transcript = html.escape(st.session_state.last_transcript)
        st.markdown(f'<div class="transcript-box">{transcript}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="summary-line">{html.escape(summary_text())}</div>', unsafe_allow_html=True)

    # Use bordered containers as card-style UI sections.
    with st.container(border=True):
        st.markdown('<div class="subheading" style="margin-top:0.1rem;">Time-stamped Segments</div>', unsafe_allow_html=True)
        # Display segment table when time-stamped ASR output is available.
        if st.session_state.last_segments:
            df = pd.DataFrame(st.session_state.last_segments)
            ordered_columns = [col for col in ["start", "end", "speaker", "source", "backend", "text"] if col in df.columns]
            df = df[ordered_columns]
            df = df.rename(columns={"start": "Start (s)", "end": "End (s)", "speaker": "Speaker", "source": "Source", "backend": "Backend", "text": "Text"})
            st.dataframe(
                df,
                width="stretch",
                hide_index=True,
                height=280,
            )
        else:
            empty_df = pd.DataFrame(columns=["Start (s)", "End (s)", "Text"])
            st.dataframe(empty_df, width="stretch", hide_index=True, height=280)



# Main Streamlit application workflow.
# Steps: configure page, initialize state, inject CSS, then render two panels.
def main() -> None:
    # Configure page before rendering any Streamlit components.
    st.set_page_config(
        page_title=APP_TITLE,
        page_icon=str(ICON_PATH) if ICON_PATH.exists() else None,
        layout="wide",
    )
    init_state()
    inject_css()
    render_header()
    left_col, right_col = st.columns([1.05, 4.2], gap="medium")
    with left_col:
        render_left_panel()
    with right_col:
        render_right_panel()


if __name__ == "__main__":
    main()
